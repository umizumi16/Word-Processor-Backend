from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
from datetime import datetime
import uuid
from bs4 import BeautifulSoup

from models import ExportRequest

router = APIRouter()

EXPORT_DIR = "static/exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

@router.post("/html")
async def export_to_html(request: ExportRequest):
    """Export document to HTML (FR-7)"""
    from routes.document import get_document
    
    try:
        doc = await get_document(request.document_id)
    except:
        raise HTTPException(status_code=404, detail="Document not found")
    
    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{doc.title}</title>
        <style>
            body {{
                font-family: 'Calibri', Arial, sans-serif;
                max-width: 8.5in;
                margin: 0 auto;
                padding: 1in;
                background: white;
            }}
            @media print {{
                body {{ padding: 0; }}
            }}
        </style>
    </head>
    <body>
        <h1>{doc.title}</h1>
        <div>{doc.content}</div>
        <footer>
            <p><small>Created: {doc.created_at}</small></p>
            <p><small>Last modified: {doc.updated_at}</small></p>
        </footer>
    </body>
    </html>
    """
    
    filename = f"{request.document_id}_{uuid.uuid4().hex[:8]}.html"
    filepath = os.path.join(EXPORT_DIR, filename)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html_template)
    
    return {"filename": filename, "path": f"/static/exports/{filename}"}

@router.post("/pdf")
async def export_to_pdf(request: ExportRequest):
    """Export document to PDF (FR-7)"""
    from routes.document import get_document
    
    try:
        doc = await get_document(request.document_id)
    except:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # First create HTML
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: 'Calibri', Arial, sans-serif;
                margin: 1in;
            }}
            @page {{
                size: A4;
                margin: 1in;
            }}
        </style>
    </head>
    <body>
        <h1>{doc.title}</h1>
        <div>{doc.content}</div>
    </body>
    </html>
    """
    
    filename = f"{request.document_id}_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join(EXPORT_DIR, filename)
    
    try:
        # Convert HTML to PDF using pdfkit
        pdfkit.from_string(html_content, filepath)
    except Exception as e:
        # Fallback: save as HTML if PDF conversion fails
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")
    
    return {"filename": filename, "path": f"/static/exports/{filename}"}

@router.post("/docx")
async def export_to_docx(request: ExportRequest):
    """Export document to DOCX (FR-7)"""
    from routes.document import get_document
    
    try:
        doc_data = await get_document(request.document_id)
    except:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Create DOCX document
    document = Document()
    
    # Add title
    title = document.add_heading(doc_data.title, 0)
    
    # Parse HTML content and convert to DOCX
    soup = BeautifulSoup(doc_data.content, 'html.parser')
    
    for element in soup.descendants:
        if element.name == 'p':
            p = document.add_paragraph(element.get_text())
        elif element.name == 'h1':
            document.add_heading(element.get_text(), 1)
        elif element.name == 'h2':
            document.add_heading(element.get_text(), 2)
        elif element.name == 'h3':
            document.add_heading(element.get_text(), 3)
        elif element.name == 'strong' or element.name == 'b':
            p = document.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
        elif element.name == 'em' or element.name == 'i':
            p = document.add_paragraph()
            run = p.add_run(element.get_text())
            run.italic = True
        elif element.name == 'u':
            p = document.add_paragraph()
            run = p.add_run(element.get_text())
            run.underline = True
    
    filename = f"{request.document_id}_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    
    document.save(filepath)
    
    return {"filename": filename, "path": f"/static/exports/{filename}"}

@router.get("/download/{filename}")
async def download_file(filename: str):
    """Download exported file"""
    filepath = os.path.join(EXPORT_DIR, filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        filepath,
        media_type='application/octet-stream',
        filename=filename
    )

@router.post("/print-preview")
async def generate_print_preview(request: ExportRequest):
    """Generate print preview HTML (FR-7)"""
    from routes.document import get_document
    
    try:
        doc = await get_document(request.document_id)
    except:
        raise HTTPException(status_code=404, detail="Document not found")
    
    print_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Print Preview - {doc.title}</title>
        <style>
            @page {{
                size: A4;
                margin: 1in;
            }}
            body {{
                font-family: 'Calibri', Arial, sans-serif;
                width: 8.5in;
                margin: 0 auto;
                padding: 1in;
                background: white;
            }}
            .page {{
                page-break-after: always;
                min-height: 11in;
                box-shadow: 0 0 10px rgba(0,0,0,0.1);
                background: white;
                padding: 1in;
                margin-bottom: 20px;
            }}
            @media print {{
                .page {{
                    box-shadow: none;
                    margin: 0;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="page">
            <h1>{doc.title}</h1>
            <div>{doc.content}</div>
        </div>
    </body>
    </html>
    """
    
    return {"html": print_html}
